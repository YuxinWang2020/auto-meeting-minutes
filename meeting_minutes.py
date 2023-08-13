import openai
import json
from os import path
from docx import Document


openai.api_key_path = 'openai_api_key.txt'
chat_model = 'gpt-3.5-turbo'
# chat_model = 'gpt-4'


def transcribe_audio(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        transcription = openai.Audio.transcribe("whisper-1", audio_file)
    return transcription['text']


def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }


def abstract_summary_extraction(transcription):
    response = openai.ChatCompletion.create(
        model=chat_model,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def key_points_extraction(transcription):
    response = openai.ChatCompletion.create(
        model=chat_model,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def action_item_extraction(transcription):
    response = openai.ChatCompletion.create(
        model=chat_model,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def sentiment_analysis(transcription):
    response = openai.ChatCompletion.create(
        model=chat_model,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)


def __load_cache(audio_file_path):
    audio_name = path.basename(audio_file_path)
    if path.exists(f".cache/{audio_name}.json"):
        with open(f".cache/{audio_name}.json", "r") as fp:
            return json.load(fp)


def __save_cache(audio_file_path, transcription = None, minutes = None):
    audio_name = path.basename(audio_file_path)
    with open(f".cache/{audio_name}.json", "w") as fp:
        json.dump({"transcription": transcription, "minutes": minutes}, fp)


if __name__ == "__main__":
    # The Whisper API only supports files that are less than 25 MB. Ref: https://platform.openai.com/docs/guides/speech-to-text/longer-inputs
    audio_file_path = "whisper test.mp3"
    cached = __load_cache(audio_file_path)
    transcription = cached['transcription'] if cached else None
    minutes = cached['minutes'] if cached else None
    if not transcription:
        transcription = transcribe_audio(audio_file_path)
        __save_cache(audio_file_path, transcription)
    if not minutes:
        minutes = meeting_minutes(transcription)
        __save_cache(audio_file_path, transcription, minutes)
    print(minutes)

    save_as_docx(minutes, 'meeting_minutes.docx')
