import openai
import json
import sys
from os import path
from docx import Document

import os
import requests as r

from dotenv import load_dotenv

load_dotenv()

chat_model = 'gpt-3.5-turbo'

MAX_TOKENS = 1000

if "http_proxy" in os.environ:
    del os.environ["http_proxy"]

def OpenAiGPTQuery(prompt, query_text):
    response = openai.ChatCompletion.create(
        model=chat_model,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": prompt
            },
            {
                "role": "user",
                "content": query_text
            }
        ]
    )
    return response['choices'][0]['message']['content']

def transcribe_audio_locally(audio_file_path, model="base"):
    import whisper
    model = whisper.load_model(model)
    result = model.transcribe(audio_file_path)
    return result["text"]

def transcribe_for_test():
    return ""

def transcribe_audio(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        transcription = openai.Audio.transcribe("whisper-1", audio_file)
    return transcription['text']


def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }


def abstract_summary_extraction(transcription):
    prompt = "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following meeting transcript and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
    return OpenAiGPTQuery(prompt, transcription)


def key_points_extraction(transcription):
    prompt = "You are a proficient AI with a specialty in distilling information into key points. Based on the following meeting transcript, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
    return OpenAiGPTQuery(prompt, transcription)


def action_item_extraction(transcription):
    prompt = "You are an AI expert in analyzing conversations and extracting action items. Please review the meeting transcript and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
    return OpenAiGPTQuery(prompt, transcription)


def sentiment_analysis(transcription):
    prompt = "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
    return OpenAiGPTQuery(prompt, transcription)

def compress_text(text, max_tokens=MAX_TOKENS):
    prompt = f"Please summarize the key details from the following meeting transcript in a comprehensive yet concise overview that is {max_tokens} words or less. Include all important information needed to thoroughly understand the key discussion points, decisions, action items, and outcomes from the meeting. Ensure the summary accurately represents the original content and context of the full transcript. Use clear, succinct language to convey the details precisely and efficiently within the word limit."
    return OpenAiGPTQuery(prompt, text)


def split_to_sentances(transcription: str) -> list:
    transcription_rm_shift = transcription.replace(">>", "")
    transcription_rm_shift = transcription_rm_shift.replace(".", ".\n")
    transcription_rm_shift = transcription_rm_shift.replace("?", "?\n")
    return [s.strip() + " " for s in transcription_rm_shift.split("\n") if len(s.strip()) > 0]

def count_tokens(text):
    return text.count(' ') + 1

def split_transcription(transcription):
    sentances = split_to_sentances(transcription)
    grouped_sentances = []

    token_count = 0
    one_group = ""
    for sentance in sentances:
        if token_count + count_tokens(sentance) > MAX_TOKENS and token_count > 0:
            grouped_sentances.append(one_group)
            token_count = 0
            one_group = ""
        token_count += count_tokens(sentance)
        one_group += sentance
    grouped_sentances.append(one_group)
    return grouped_sentances

def split_and_compress(transcription):
    compression_rate = MAX_TOKENS / count_tokens(transcription)
    splited_transcriptions = split_transcription(transcription)
    compressed_texts = [compress_text(splited, int(count_tokens(splited) * compression_rate)) for splited in splited_transcriptions]
    return " ".join(compressed_texts)

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)


def __load_cache(audio_file_path):
    # audio_name = path.basename(audio_file_path)
    # if path.exists(f".cache/{audio_name}.json"):
    #     with open(f".cache/{audio_name}.json", "r") as fp:
    #         return json.load(fp)
    return None

def __save_cache(audio_file_path, transcription = None, minutes = None):
    audio_name = path.basename(audio_file_path)
    with open(f".cache/{audio_name}.json", "w") as fp:
        json.dump({"transcription": transcription, "minutes": minutes}, fp)


if __name__ == "__main__":
    if len(sys.argv) >= 2:
        audio_file_path = sys.argv[1]
    else:
        audio_file_path = "audio"
    cached = __load_cache(audio_file_path)
    transcription = cached['transcription'] if cached else None
    minutes = cached['minutes'] if cached else None
    if not transcription:
        transcription = transcribe_for_test()
        __save_cache(audio_file_path, transcription)

    transcription = split_and_compress(transcription)

    if not minutes:
        minutes = meeting_minutes(transcription)
        __save_cache(audio_file_path, transcription, minutes)
    print(minutes)

    save_as_docx(minutes, 'meeting_minutes.docx')
